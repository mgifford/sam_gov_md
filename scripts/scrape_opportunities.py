#!/usr/bin/env python3
"""Download and extract text from PDF and Word document attachments linked in SAM.gov opportunities."""

from __future__ import annotations

import argparse
import csv
import html as html_module
import io
import json
import os
import re
import time
import unicodedata
from pathlib import Path
from typing import Optional
from urllib.parse import urljoin, urlparse

import docx
import pdfplumber
import requests

DEFAULT_CSV = "data/ContractOpportunitiesFullCSV.csv"
DEFAULT_OUTPUT_DIR = "docs/opportunities"
DEFAULT_LIMIT = 50
DEFAULT_TIMEOUT = 30
SAM_GOV_API_BASE = "https://api.sam.gov/opportunities/v1"

_DOCX_CONTENT_TYPES = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
)


def is_pdf_url(url: str) -> bool:
    parsed = urlparse(url)
    return parsed.path.lower().endswith(".pdf")


def is_docx_url(url: str) -> bool:
    parsed = urlparse(url)
    return parsed.path.lower().endswith((".docx", ".doc"))


def fetch_sam_gov_attachments(
    notice_id: str,
    api_key: str = "DEMO_KEY",
    timeout: int = DEFAULT_TIMEOUT,
) -> list[dict]:
    """Fetch attachment URLs for an opportunity via the SAM.gov public API.

    Returns a list of dicts with ``url`` and ``filename`` keys.  Falls back to
    an empty list on any error so the caller can continue without crashing.
    """
    url = f"{SAM_GOV_API_BASE}/search?noticeid={notice_id}&api_key={api_key}"
    try:
        resp = requests.get(
            url,
            timeout=timeout,
            headers={"User-Agent": "Mozilla/5.0 (compatible; SAMGovBot/1.0)"},
        )
        if resp.status_code != 200:
            return []
        data = resp.json()

        # The SAM.gov API may use different envelope structures across versions.
        # Walk through several known shapes to find the opportunity record.
        results: list[dict] = (
            data.get("opportunitiesData")
            or data.get("_embedded", {}).get("results", [])
            or []
        )
        if not results:
            return []

        opp = results[0]
        resource_links: list[str] = opp.get("resourceLinks") or []
        attachments = []
        for link in resource_links:
            if not link:
                continue
            filename = Path(urlparse(link).path).name or "attachment"
            attachments.append({"url": link, "filename": filename})
        return attachments
    except Exception:
        return []


def fetch_html_summary_and_pdfs(
    url: str,
    timeout: int = DEFAULT_TIMEOUT,
) -> tuple[str, list[str]]:
    """Fetch an HTML URL and return (summary_text, list_of_pdf_urls).

    Uses BeautifulSoup when available; falls back to a plain-text excerpt if
    the import fails so that the caller still gets something useful.
    """
    try:
        resp = requests.get(
            url,
            timeout=timeout,
            headers={"User-Agent": "Mozilla/5.0 (compatible; SAMGovBot/1.0)"},
            allow_redirects=True,
        )
        resp.raise_for_status()
        content_type = resp.headers.get("Content-Type", "").lower()
        # If the server actually returned a PDF, signal that to the caller.
        if "pdf" in content_type:
            return "", [url]

        try:
            from bs4 import BeautifulSoup  # optional dependency

            soup = BeautifulSoup(resp.content, "lxml")
            # Remove noise tags
            for tag in soup(["script", "style", "nav", "header", "footer", "noscript"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
            # Limit to a reasonable summary size
            text = "\n".join(line for line in text.splitlines() if line.strip())
            text = text[:4000]

            # Find embedded PDF links
            pdf_links: list[str] = []
            for a in soup.find_all("a", href=True):
                href = str(a["href"]).strip()
                if not href or href.startswith("#"):
                    continue
                if ".pdf" in href.lower():
                    full = href if href.startswith("http") else urljoin(url, href)
                    if full not in pdf_links:
                        pdf_links.append(full)

            return text, pdf_links

        except ImportError:
            # BeautifulSoup not installed – return a truncated raw-text excerpt
            raw = resp.text[:4000]
            # Strip tags with a simple regex for a rough plain-text view
            plain = re.sub(r"<[^>]+>", " ", raw)
            plain = re.sub(r"\s+", " ", plain).strip()
            return plain[:4000], []

    except Exception as exc:
        return f"[Link fetch error: {exc}]", []


def fetch_document_bytes(
    url: str, timeout: int = DEFAULT_TIMEOUT, retries: int = 2
) -> tuple[Optional[bytes], Optional[str]]:
    """Download a document from *url* and return ``(bytes, doc_type)``.

    *doc_type* is ``"pdf"`` or ``"docx"``.  Returns ``(None, None)`` when the
    resource cannot be fetched or is not a recognised document type.
    """
    last_err: Optional[Exception] = None
    for attempt in range(1, retries + 1):
        try:
            resp = requests.get(
                url,
                timeout=timeout,
                headers={"User-Agent": "Mozilla/5.0 (compatible; SAMGovBot/1.0)"},
                allow_redirects=True,
            )
            resp.raise_for_status()
            content_type = resp.headers.get("Content-Type", "").lower()
            if "pdf" in content_type or is_pdf_url(url):
                return resp.content, "pdf"
            if any(content_type.startswith(ct) for ct in _DOCX_CONTENT_TYPES) or is_docx_url(url):
                return resp.content, "docx"
            # Not a recognised document type
            return None, None
        except Exception as exc:
            last_err = exc
            if attempt < retries:
                time.sleep(2 ** attempt)
    return None, None


def clean_document_text(text: str) -> str:
    """Normalize text extracted from documents: fix common encoding artifacts."""
    if not text:
        return text
    # Unescape any HTML entities that may appear in the source
    text = html_module.unescape(text)
    # Normalize Unicode to NFC (composed form) to consolidate combining chars
    text = unicodedata.normalize("NFC", text)
    return text


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            pages: list[str] = []
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(text)
            raw = "\n\n".join(pages)
            return clean_document_text(raw)
    except Exception as exc:
        return f"[PDF extraction error: {exc}]"


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract plain text from a Word document (.docx/.doc)."""
    try:
        document = docx.Document(io.BytesIO(docx_bytes))
        parts: list[str] = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        # Extract text from tables (e.g. pricing or line-item tables)
        for table in document.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    parts.append(" | ".join(row_texts))
        raw = "\n\n".join(parts)
        return clean_document_text(raw)
    except Exception as exc:
        return f"[Word document extraction error: {exc}]"


def write_opportunity_pdf_content(
    notice_id: str,
    title: str,
    attachments: list[dict],
    output_dir: Path,
    save_pdf: bool = False,
) -> None:
    """Write ``pdf_content.md`` for an opportunity.

    ``attachments`` is a list of dicts with the following keys:

    - ``url``       – source URL (PDF, Word document, or HTML)
    - ``filename``  – display filename / label
    - ``text``      – extracted text content
    - ``pdf_bytes`` – raw PDF bytes (optional, used when ``save_pdf=True``)
    - ``doc_bytes`` – raw Word document bytes (optional, used when ``save_pdf=True``)
    - ``kind``      – ``"pdf"``, ``"docx"``, ``"html"``, or ``"none"``
    """
    doc_dir = output_dir / notice_id
    doc_dir.mkdir(parents=True, exist_ok=True)

    lines: list[str] = [
        f"# Attachments: {title}",
        "",
        f"- Notice ID: {notice_id}",
        "",
    ]

    if not attachments:
        lines += [
            "_No PDF attachments or document links were found for this opportunity._",
            "",
            f"- SAM.gov opportunity page: https://sam.gov/workspace/contract/opp/{notice_id}/view",
        ]
    else:
        for i, att in enumerate(attachments, 1):
            url = att.get("url", "")
            filename = att.get("filename", f"Attachment {i}")
            text = att.get("text", "")
            kind = att.get("kind", "pdf")

            heading = f"## Attachment {i}: {filename}"
            lines.append(heading)
            lines.append("")
            if url:
                lines.append(f"- URL: {url}")
            lines.append("")

            if kind == "html":
                lines.append("### Page Summary")
            else:
                lines.append("### Extracted Text")
            lines.append("")
            lines.append(text or "_No text could be extracted._")
            lines.append("")
            lines.append("---")
            lines.append("")

            # Optionally persist the raw attachment file next to the markdown
            if save_pdf and kind == "pdf" and att.get("pdf_bytes"):
                raw_bytes = att["pdf_bytes"]
                url_path = urlparse(url).path
                raw_name = Path(url_path).name or "attachment"
                raw_stem = Path(raw_name).stem
                safe_stem = re.sub(r"[^A-Za-z0-9_-]", "_", raw_stem) or "attachment"
                safe_name = f"{safe_stem}.pdf"
                (doc_dir / safe_name).write_bytes(raw_bytes)
            elif save_pdf and kind == "docx" and att.get("doc_bytes"):
                raw_bytes = att["doc_bytes"]
                url_path = urlparse(url).path
                raw_name = Path(url_path).name or "attachment"
                raw_stem = Path(raw_name).stem
                url_suffix = Path(url_path).suffix.lower()
                safe_ext = url_suffix if url_suffix in (".docx", ".doc") else ".docx"
                safe_stem = re.sub(r"[^A-Za-z0-9_-]", "_", raw_stem) or "attachment"
                safe_name = f"{safe_stem}{safe_ext}"
                (doc_dir / safe_name).write_bytes(raw_bytes)

    (doc_dir / "pdf_content.md").write_text("\n".join(lines), encoding="utf-8")


# ---------------------------------------------------------------------------
# Legacy single-PDF helper kept for backwards compatibility with any callers
# that still use the old three-positional-arg signature.
# ---------------------------------------------------------------------------
def _write_opportunity_pdf_content_legacy(
    notice_id: str,
    title: str,
    pdf_url: str,
    text: str,
    output_dir: Path,
    pdf_bytes: Optional[bytes] = None,
    save_pdf: bool = False,
) -> None:
    att = {
        "url": pdf_url,
        "filename": Path(urlparse(pdf_url).path).name or "attachment.pdf",
        "text": text,
        "kind": "pdf",
        "pdf_bytes": pdf_bytes,
    }
    write_opportunity_pdf_content(notice_id, title, [att], output_dir, save_pdf=save_pdf)


def write_opportunity_docx_content(
    notice_id: str,
    title: str,
    doc_url: str,
    text: str,
    output_dir: Path,
    doc_bytes: Optional[bytes] = None,
    save_doc: bool = False,
) -> None:
    """Write extracted Word document text to ``docx_content.md``."""
    doc_dir = output_dir / notice_id
    doc_dir.mkdir(parents=True, exist_ok=True)
    lines = [
        f"# Word Document Content: {title}",
        "",
        f"- Notice ID: {notice_id}",
        f"- Document URL: {doc_url}",
        "",
        "## Extracted Text",
        "",
        text or "_No text could be extracted from this Word document._",
    ]
    (doc_dir / "docx_content.md").write_text("\n".join(lines), encoding="utf-8")

    if save_doc and doc_bytes:
        url_path = urlparse(doc_url).path
        raw_name = Path(url_path).name or "attachment"
        raw_stem = Path(raw_name).stem
        # Preserve the original extension (.docx or .doc) when saving the raw file
        url_suffix = Path(url_path).suffix.lower()
        safe_ext = url_suffix if url_suffix in (".docx", ".doc") else ".docx"
        safe_stem = re.sub(r"[^A-Za-z0-9_-]", "_", raw_stem) or "attachment"
        safe_name = f"{safe_stem}{safe_ext}"
        (doc_dir / safe_name).write_bytes(doc_bytes)


def load_csv(csv_path: str) -> list[dict]:
    rows: list[dict] = []
    try:
        # SAM.gov CSVs commonly use Windows-1252 encoding; errors="replace" handles any
        # remaining invalid bytes without crashing on unusual field values.
        with open(csv_path, "r", encoding="windows-1252", errors="replace") as f:
            reader = csv.DictReader(f)
            for row in reader:
                rows.append(row)
    except Exception as exc:
        print(f"Error reading CSV {csv_path}: {exc}")
    return rows


def select_candidates(rows: list[dict], limit: int) -> list[dict]:
    """Select top candidates for attachment/PDF extraction.

    Rows with a non-empty ``AdditionalInfoLink`` are prioritised because they
    already have a direct link to process.  All other rows are included
    afterwards so that the SAM.gov API can be queried for their attachments.
    """
    with_links = [
        row for row in rows if (row.get("AdditionalInfoLink") or "").strip()
    ]
    without_links = [
        row for row in rows if not (row.get("AdditionalInfoLink") or "").strip()
    ]
    # Sort each group by posted date descending (most recent first)
    for group in (with_links, without_links):
        group.sort(key=lambda r: (r.get("PostedDate") or ""), reverse=True)
    combined = with_links + without_links
    return combined[:limit]


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Download and extract PDF and Word document attachments from SAM.gov opportunities"
    )
    parser.add_argument(
        "--csv",
        default=DEFAULT_CSV,
        help=f"Path to the full CSV extract (default: {DEFAULT_CSV})",
    )
    parser.add_argument(
        "--limit",
        type=int,
        default=DEFAULT_LIMIT,
        help=f"Max number of opportunities to process (default: {DEFAULT_LIMIT})",
    )
    parser.add_argument(
        "--output-dir",
        default=DEFAULT_OUTPUT_DIR,
        help=f"Output directory for extracted content (default: {DEFAULT_OUTPUT_DIR})",
    )
    parser.add_argument(
        "--timeout",
        type=int,
        default=DEFAULT_TIMEOUT,
        help=f"HTTP request timeout in seconds (default: {DEFAULT_TIMEOUT})",
    )
    parser.add_argument(
        "--summary-output",
        default="data/today/pdf_extraction_summary.json",
        help="Path to write extraction summary JSON",
    )
    parser.add_argument(
        "--save-pdf",
        action="store_true",
        help=(
            "Save the raw attachment file alongside the extracted markdown "
            "in the opportunity directory (applies to both PDFs and Word documents)"
        ),
    )
    parser.add_argument(
        "--max-file-size",
        type=int,
        default=50 * 1024 * 1024,  # 50 MB
        help="Maximum attachment file size in bytes to save (default: 52428800 = 50 MB)",
    )
    parser.add_argument(
        "--api-key",
        default=os.environ.get("SAM_API_KEY", "DEMO_KEY"),
        help=(
            "SAM.gov API key used to fetch attachment lists. "
            "Defaults to the SAM_API_KEY environment variable or 'DEMO_KEY'."
        ),
    )
    args = parser.parse_args()

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"Loading CSV from {args.csv}")
    rows = load_csv(args.csv)
    if not rows:
        print("No rows loaded from CSV. Exiting.")
        return

    print(f"Total rows in CSV: {len(rows)}")
    candidates = select_candidates(rows, args.limit)
    print(f"Candidates to process: {len(candidates)}")

    extracted = 0
    skipped = 0
    errors = 0
    summary_records: list[dict] = []

    for i, row in enumerate(candidates, start=1):
        notice_id = (row.get("NoticeId") or "").strip()
        title = (row.get("Title") or "Untitled").strip()
        additional_link = (row.get("AdditionalInfoLink") or "").strip()

        if not notice_id:
            skipped += 1
            continue

        print(f"  [{i}/{len(candidates)}] {notice_id}: {title[:60]}")

        # Collect all attachments for this opportunity
        attachments: list[dict] = []

        # ── Step 1: process AdditionalInfoLink if present ──────────────────
        if additional_link:
            if is_pdf_url(additional_link) or is_docx_url(additional_link):
                # Direct document URL (PDF or Word)
                doc_bytes, doc_type = fetch_document_bytes(additional_link, timeout=args.timeout)
                if doc_bytes is not None:
                    doc_too_large = args.save_pdf and len(doc_bytes) > args.max_file_size
                    if doc_type == "docx":
                        text = extract_text_from_docx(doc_bytes)
                        filename = Path(urlparse(additional_link).path).name or "attachment.docx"
                        attachments.append({
                            "url": additional_link,
                            "filename": filename,
                            "text": text,
                            "kind": "docx",
                            "doc_bytes": doc_bytes if not doc_too_large else None,
                        })
                        print(f"    AdditionalInfoLink Word doc: ~{len(text.split())} words")
                    else:
                        text = extract_text_from_pdf(doc_bytes)
                        filename = Path(urlparse(additional_link).path).name or "attachment.pdf"
                        attachments.append({
                            "url": additional_link,
                            "filename": filename,
                            "text": text,
                            "kind": "pdf",
                            "pdf_bytes": doc_bytes if not doc_too_large else None,
                        })
                        print(f"    AdditionalInfoLink PDF: ~{len(text.split())} words")
                else:
                    print(f"    AdditionalInfoLink document download failed: {additional_link[:60]}")
            else:
                # HTML / portal link – summarise the page and look for embedded PDFs
                html_text, embedded_pdfs = fetch_html_summary_and_pdfs(
                    additional_link, timeout=args.timeout
                )
                if html_text:
                    attachments.append({
                        "url": additional_link,
                        "filename": f"Link: {additional_link[:60]}",
                        "text": html_text,
                        "kind": "html",
                    })
                    print(f"    AdditionalInfoLink HTML summary: ~{len(html_text.split())} words")
                for pdf_url in embedded_pdfs:
                    doc_bytes, doc_type = fetch_document_bytes(pdf_url, timeout=args.timeout)
                    if doc_bytes is not None:
                        doc_too_large = args.save_pdf and len(doc_bytes) > args.max_file_size
                        if doc_type == "docx":
                            text = extract_text_from_docx(doc_bytes)
                            filename = Path(urlparse(pdf_url).path).name or "attachment.docx"
                            attachments.append({
                                "url": pdf_url,
                                "filename": filename,
                                "text": text,
                                "kind": "docx",
                                "doc_bytes": doc_bytes if not doc_too_large else None,
                            })
                            print(f"    Embedded Word doc: {filename} (~{len(text.split())} words)")
                        else:
                            text = extract_text_from_pdf(doc_bytes)
                            filename = Path(urlparse(pdf_url).path).name or "attachment.pdf"
                            attachments.append({
                                "url": pdf_url,
                                "filename": filename,
                                "text": text,
                                "kind": "pdf",
                                "pdf_bytes": doc_bytes if not doc_too_large else None,
                            })
                            print(f"    Embedded PDF: {filename} (~{len(text.split())} words)")

        # ── Step 2: query SAM.gov API for additional/all attachments ───────
        api_attachments = fetch_sam_gov_attachments(
            notice_id, api_key=args.api_key, timeout=args.timeout
        )
        seen_urls = {a["url"] for a in attachments}
        for att_meta in api_attachments:
            att_url = att_meta.get("url", "")
            att_filename = att_meta.get("filename", "attachment")
            if not att_url or att_url in seen_urls:
                continue
            seen_urls.add(att_url)
            if is_pdf_url(att_url) or is_docx_url(att_url):
                doc_bytes, doc_type = fetch_document_bytes(att_url, timeout=args.timeout)
                if doc_bytes is not None:
                    doc_too_large = args.save_pdf and len(doc_bytes) > args.max_file_size
                    if doc_type == "docx":
                        text = extract_text_from_docx(doc_bytes)
                        attachments.append({
                            "url": att_url,
                            "filename": att_filename,
                            "text": text,
                            "kind": "docx",
                            "doc_bytes": doc_bytes if not doc_too_large else None,
                        })
                        print(f"    API Word doc: {att_filename} (~{len(text.split())} words)")
                    else:
                        text = extract_text_from_pdf(doc_bytes)
                        attachments.append({
                            "url": att_url,
                            "filename": att_filename,
                            "text": text,
                            "kind": "pdf",
                            "pdf_bytes": doc_bytes if not doc_too_large else None,
                        })
                        print(f"    API PDF: {att_filename} (~{len(text.split())} words)")
            else:
                html_text, embedded_pdfs = fetch_html_summary_and_pdfs(
                    att_url, timeout=args.timeout
                )
                if html_text:
                    attachments.append({
                        "url": att_url,
                        "filename": att_filename,
                        "text": html_text,
                        "kind": "html",
                    })
                    print(f"    API link summary: {att_filename}")
                for pdf_url in embedded_pdfs:
                    if pdf_url in seen_urls:
                        continue
                    seen_urls.add(pdf_url)
                    doc_bytes, doc_type = fetch_document_bytes(pdf_url, timeout=args.timeout)
                    if doc_bytes is not None:
                        doc_too_large = args.save_pdf and len(doc_bytes) > args.max_file_size
                        if doc_type == "docx":
                            text = extract_text_from_docx(doc_bytes)
                            filename = Path(urlparse(pdf_url).path).name or "attachment.docx"
                            attachments.append({
                                "url": pdf_url,
                                "filename": filename,
                                "text": text,
                                "kind": "docx",
                                "doc_bytes": doc_bytes if not doc_too_large else None,
                            })
                            print(f"    API embedded Word doc: {filename}")
                        else:
                            text = extract_text_from_pdf(doc_bytes)
                            filename = Path(urlparse(pdf_url).path).name or "attachment.pdf"
                            attachments.append({
                                "url": pdf_url,
                                "filename": filename,
                                "text": text,
                                "kind": "pdf",
                                "pdf_bytes": doc_bytes if not doc_too_large else None,
                            })
                            print(f"    API embedded PDF: {filename}")

        # ── Step 3: write the combined pdf_content.md ───────────────────────
        try:
            write_opportunity_pdf_content(
                notice_id,
                title,
                attachments,
                output_dir,
                save_pdf=args.save_pdf,
            )
            if attachments:
                total_words = sum(len(a.get("text", "").split()) for a in attachments)
                print(f"    Wrote pdf_content.md ({len(attachments)} attachment(s), ~{total_words} words total)")
                extracted += 1
                summary_records.append({
                    "notice_id": notice_id,
                    "status": "ok",
                    "attachments": len(attachments),
                    "words": total_words,
                })
            else:
                print(f"    No attachments found – wrote 'no attachments' note")
                skipped += 1
                summary_records.append({
                    "notice_id": notice_id,
                    "status": "no_attachments",
                })
        except Exception as exc:
            print(f"    Error writing pdf_content.md: {exc}")
            errors += 1
            summary_records.append(
                {"notice_id": notice_id, "status": "error", "error": str(exc)}
            )

    summary = {
        "csv": args.csv,
        "limit": args.limit,
        "total_rows": len(rows),
        "candidates": len(candidates),
        "extracted": extracted,
        "skipped": skipped,
        "errors": errors,
        "records": summary_records,
    }

    summary_path = Path(args.summary_output)
    summary_path.parent.mkdir(parents=True, exist_ok=True)
    summary_path.write_text(json.dumps(summary, indent=2), encoding="utf-8")

    print(f"\nExtraction complete: {extracted} with attachments, {skipped} no attachments, {errors} errors")
    print(f"Summary written to {summary_path}")


if __name__ == "__main__":
    main()
