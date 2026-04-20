"""Tests for scrape_opportunities.py."""

from __future__ import annotations

import csv
import io
import json
import os
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

import scrape_opportunities as so


# ---------------------------------------------------------------------------
# is_pdf_url
# ---------------------------------------------------------------------------

class TestIsPdfUrl:
    def test_simple_pdf_extension(self) -> None:
        assert so.is_pdf_url("https://example.com/document.pdf") is True

    def test_uppercase_extension(self) -> None:
        assert so.is_pdf_url("https://example.com/document.PDF") is True

    def test_mixed_case_extension(self) -> None:
        assert so.is_pdf_url("https://example.com/document.Pdf") is True

    def test_pdf_with_query_string(self) -> None:
        # Query string is not part of the path, but the path still ends in .pdf
        assert so.is_pdf_url("https://example.com/file.pdf?v=1") is True

    def test_non_pdf_url(self) -> None:
        assert so.is_pdf_url("https://example.com/document.docx") is False

    def test_html_url(self) -> None:
        assert so.is_pdf_url("https://example.com/page.html") is False

    def test_empty_string(self) -> None:
        assert so.is_pdf_url("") is False

    def test_url_with_pdf_in_path_but_not_extension(self) -> None:
        assert so.is_pdf_url("https://example.com/pdf-docs/report.html") is False


# ---------------------------------------------------------------------------
# is_docx_url
# ---------------------------------------------------------------------------

class TestIsDocxUrl:
    def test_docx_extension(self) -> None:
        assert so.is_docx_url("https://example.com/document.docx") is True

    def test_doc_extension(self) -> None:
        assert so.is_docx_url("https://example.com/document.doc") is True

    def test_uppercase_extension(self) -> None:
        assert so.is_docx_url("https://example.com/document.DOCX") is True

    def test_pdf_is_not_docx(self) -> None:
        assert so.is_docx_url("https://example.com/document.pdf") is False

    def test_html_is_not_docx(self) -> None:
        assert so.is_docx_url("https://example.com/page.html") is False

    def test_empty_string(self) -> None:
        assert so.is_docx_url("") is False


# ---------------------------------------------------------------------------
# clean_document_text
# ---------------------------------------------------------------------------

class TestCleanDocumentText:
    def test_empty_string_passthrough(self) -> None:
        assert so.clean_document_text("") == ""

    def test_none_like_empty(self) -> None:
        # Function returns input unchanged if falsy
        result = so.clean_document_text("")
        assert result == ""

    def test_html_entities_unescaped(self) -> None:
        result = so.clean_document_text("AT&amp;T &lt;Corp&gt;")
        assert result == "AT&T <Corp>"

    def test_unicode_normalization_nfc(self) -> None:
        # Decomposed é (e + combining acute) should become composed é
        decomposed = "e\u0301"  # e + combining acute accent
        result = so.clean_document_text(decomposed)
        assert result == "\xe9"  # precomposed é

    def test_plain_text_unchanged(self) -> None:
        text = "Hello, world! This is a test."
        assert so.clean_document_text(text) == text

    def test_whitespace_preserved(self) -> None:
        text = "  leading and trailing  "
        assert so.clean_document_text(text) == text


# ---------------------------------------------------------------------------
# load_csv
# ---------------------------------------------------------------------------

class TestLoadCsv:
    def _write_csv(self, tmp_path: Path, rows: list[dict], fieldnames: list[str]) -> str:
        csv_file = str(tmp_path / "test.csv")
        with open(csv_file, "w", newline="", encoding="windows-1252") as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(rows)
        return csv_file

    def test_basic_load(self, tmp_path: Path) -> None:
        fields = ["NoticeId", "Title", "Agency"]
        rows = [{"NoticeId": "ABC123", "Title": "Test", "Agency": "DoD"}]
        csv_file = self._write_csv(tmp_path, rows, fields)
        result = so.load_csv(csv_file)
        assert len(result) == 1
        assert result[0]["NoticeId"] == "ABC123"

    def test_multiple_rows(self, tmp_path: Path) -> None:
        fields = ["NoticeId", "Title"]
        rows = [
            {"NoticeId": "A1", "Title": "First"},
            {"NoticeId": "A2", "Title": "Second"},
            {"NoticeId": "A3", "Title": "Third"},
        ]
        csv_file = self._write_csv(tmp_path, rows, fields)
        result = so.load_csv(csv_file)
        assert len(result) == 3

    def test_missing_file_returns_empty_list(self) -> None:
        result = so.load_csv("/nonexistent/path/file.csv")
        assert result == []

    def test_empty_csv_returns_empty_list(self, tmp_path: Path) -> None:
        csv_file = str(tmp_path / "empty.csv")
        with open(csv_file, "w", encoding="windows-1252") as f:
            f.write("NoticeId,Title\n")
        result = so.load_csv(csv_file)
        assert result == []


# ---------------------------------------------------------------------------
# select_candidates
# ---------------------------------------------------------------------------

class TestSelectCandidates:
    def _row(self, notice_id: str, date: str = "2024-01-01", link: str = "") -> dict:
        return {"NoticeId": notice_id, "PostedDate": date, "AdditionalInfoLink": link}

    def test_rows_with_links_come_first(self) -> None:
        rows = [
            self._row("NO_LINK_1", "2024-01-01"),
            self._row("WITH_LINK_1", "2024-01-01", "https://example.com/doc.pdf"),
            self._row("NO_LINK_2", "2024-01-02"),
        ]
        result = so.select_candidates(rows, limit=10)
        assert result[0]["NoticeId"] == "WITH_LINK_1"

    def test_limit_applied(self) -> None:
        rows = [self._row(f"ID{i}", link="https://link.com") for i in range(10)]
        result = so.select_candidates(rows, limit=3)
        assert len(result) == 3

    def test_sorted_by_date_descending_within_group(self) -> None:
        rows = [
            self._row("LINK_OLD", "2024-01-01", "https://link.com"),
            self._row("LINK_NEW", "2024-06-01", "https://link.com"),
            self._row("LINK_MID", "2024-03-01", "https://link.com"),
        ]
        result = so.select_candidates(rows, limit=10)
        dates = [r["PostedDate"] for r in result]
        assert dates == sorted(dates, reverse=True)

    def test_empty_rows_returns_empty(self) -> None:
        assert so.select_candidates([], limit=10) == []

    def test_all_rows_no_links(self) -> None:
        rows = [self._row(f"ID{i}", "2024-0{i+1}-01") for i in range(3)]
        result = so.select_candidates(rows, limit=10)
        assert len(result) == 3
        # All are in the "no links" group, sorted newest first
        dates = [r["PostedDate"] for r in result]
        assert dates == sorted(dates, reverse=True)


# ---------------------------------------------------------------------------
# extract_text_from_pdf (mocked – avoids pdfplumber dependency at test time)
# ---------------------------------------------------------------------------

class TestExtractTextFromPdf:
    def test_returns_error_string_on_bad_bytes(self) -> None:
        result = so.extract_text_from_pdf(b"not a valid pdf")
        assert result.startswith("[PDF extraction error:")

    def test_valid_pdf_extraction(self) -> None:
        # Mock pdfplumber to return predictable text
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Hello from PDF"
        mock_pdf = MagicMock()
        mock_pdf.pages = [mock_page]
        mock_pdf.__enter__ = MagicMock(return_value=mock_pdf)
        mock_pdf.__exit__ = MagicMock(return_value=False)

        with patch("pdfplumber.open", return_value=mock_pdf):
            result = so.extract_text_from_pdf(b"%PDF-1.4 fake")
        assert "Hello from PDF" in result

    def test_empty_pages_returns_empty_string(self) -> None:
        mock_page = MagicMock()
        mock_page.extract_text.return_value = ""
        mock_pdf = MagicMock()
        mock_pdf.pages = [mock_page]
        mock_pdf.__enter__ = MagicMock(return_value=mock_pdf)
        mock_pdf.__exit__ = MagicMock(return_value=False)

        with patch("pdfplumber.open", return_value=mock_pdf):
            result = so.extract_text_from_pdf(b"%PDF-1.4 fake")
        assert result == ""

    def test_multiple_pages_joined(self) -> None:
        mock_page1 = MagicMock()
        mock_page1.extract_text.return_value = "Page one"
        mock_page2 = MagicMock()
        mock_page2.extract_text.return_value = "Page two"
        mock_pdf = MagicMock()
        mock_pdf.pages = [mock_page1, mock_page2]
        mock_pdf.__enter__ = MagicMock(return_value=mock_pdf)
        mock_pdf.__exit__ = MagicMock(return_value=False)

        with patch("pdfplumber.open", return_value=mock_pdf):
            result = so.extract_text_from_pdf(b"%PDF-1.4 fake")
        assert "Page one" in result
        assert "Page two" in result


# ---------------------------------------------------------------------------
# extract_text_from_docx
# ---------------------------------------------------------------------------

class TestExtractTextFromDocx:
    def test_returns_error_string_on_bad_bytes(self) -> None:
        result = so.extract_text_from_docx(b"not a valid docx")
        assert result.startswith("[Word document extraction error:")

    def test_valid_docx_extraction(self) -> None:
        import docx as python_docx

        # Build a minimal valid DOCX in memory
        doc = python_docx.Document()
        doc.add_paragraph("Hello from DOCX")
        buf = io.BytesIO()
        doc.save(buf)
        result = so.extract_text_from_docx(buf.getvalue())
        assert "Hello from DOCX" in result

    def test_table_content_extracted(self) -> None:
        import docx as python_docx

        doc = python_docx.Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Cell A"
        table.cell(0, 1).text = "Cell B"
        buf = io.BytesIO()
        doc.save(buf)
        result = so.extract_text_from_docx(buf.getvalue())
        assert "Cell A" in result
        assert "Cell B" in result

    def test_empty_docx_returns_empty_string(self) -> None:
        import docx as python_docx

        doc = python_docx.Document()
        buf = io.BytesIO()
        doc.save(buf)
        result = so.extract_text_from_docx(buf.getvalue())
        # A blank document with no content should produce empty or whitespace-only string
        assert result.strip() == ""


# ---------------------------------------------------------------------------
# fetch_sam_gov_attachments (mocked)
# ---------------------------------------------------------------------------

class TestFetchSamGovAttachments:
    def test_returns_empty_list_on_http_error(self) -> None:
        mock_resp = MagicMock()
        mock_resp.status_code = 404
        with patch("requests.get", return_value=mock_resp):
            result = so.fetch_sam_gov_attachments("NOTICE123")
        assert result == []

    def test_returns_empty_list_on_exception(self) -> None:
        with patch("requests.get", side_effect=Exception("network error")):
            result = so.fetch_sam_gov_attachments("NOTICE123")
        assert result == []

    def test_returns_empty_list_when_no_results(self) -> None:
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {"opportunitiesData": []}
        with patch("requests.get", return_value=mock_resp):
            result = so.fetch_sam_gov_attachments("NOTICE123")
        assert result == []

    def test_parses_resource_links(self) -> None:
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "opportunitiesData": [
                {"resourceLinks": ["https://example.com/files/statement.pdf"]}
            ]
        }
        with patch("requests.get", return_value=mock_resp):
            result = so.fetch_sam_gov_attachments("NOTICE123")
        assert len(result) == 1
        assert result[0]["url"] == "https://example.com/files/statement.pdf"
        assert result[0]["filename"] == "statement.pdf"

    def test_handles_embedded_results_envelope(self) -> None:
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "_embedded": {
                "results": [
                    {"resourceLinks": ["https://example.com/spec.docx"]}
                ]
            }
        }
        with patch("requests.get", return_value=mock_resp):
            result = so.fetch_sam_gov_attachments("NOTICE456")
        assert len(result) == 1
        assert result[0]["filename"] == "spec.docx"

    # ------------------------------------------------------------------
    # Tests for the primary noticedesc endpoint / fileInformation shape
    # ------------------------------------------------------------------

    def test_parses_file_information_attachments(self) -> None:
        """Primary path: noticedesc response with fileInformation.fileID."""
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "attachments": [
                {
                    "name": "RFQ Letter.pdf",
                    "fileInformation": {
                        "fileID": "abc-123",
                        "fileName": "RFQ Letter.pdf",
                    },
                }
            ]
        }
        with patch("requests.get", return_value=mock_resp):
            result = so.fetch_sam_gov_attachments("NOTICE789", api_key="MYKEY")
        assert len(result) == 1
        assert result[0]["filename"] == "RFQ Letter.pdf"
        assert "NOTICE789" in result[0]["url"]
        assert "abc-123" in result[0]["url"]
        assert "MYKEY" in result[0]["url"]

    def test_download_url_format(self) -> None:
        """Constructed download URL matches expected SAM.gov v3 path."""
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "attachments": [
                {
                    "fileInformation": {
                        "fileID": "file-uuid-001",
                        "fileName": "SOW.pdf",
                    }
                }
            ]
        }
        with patch("requests.get", return_value=mock_resp):
            result = so.fetch_sam_gov_attachments("NOTICEID", api_key="TESTKEY")
        assert len(result) == 1
        expected_url = (
            "https://sam.gov/api/prod/opps/v3/opportunities/NOTICEID"
            "/resources/files/file-uuid-001/download?api_key=TESTKEY"
        )
        assert result[0]["url"] == expected_url

    def test_parses_file_infomation_typo_variant(self) -> None:
        """Fallback spelling 'fileInfomation' (missing 'r') is also accepted."""
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "attachments": [
                {
                    "name": "attachment.pdf",
                    "fileInfomation": {
                        "fileID": "typo-variant-id",
                        "fileName": "attachment.pdf",
                    },
                }
            ]
        }
        with patch("requests.get", return_value=mock_resp):
            result = so.fetch_sam_gov_attachments("NOTICE_T", api_key="K")
        assert len(result) == 1
        assert "typo-variant-id" in result[0]["url"]

    def test_attachment_without_file_id_is_skipped(self) -> None:
        """Attachment objects that have no fileID are skipped."""
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "attachments": [
                {"name": "no_id.pdf", "fileInformation": {}},
                {
                    "name": "has_id.pdf",
                    "fileInformation": {"fileID": "good-id", "fileName": "has_id.pdf"},
                },
            ]
        }
        with patch("requests.get", return_value=mock_resp):
            result = so.fetch_sam_gov_attachments("NOTICE_S")
        assert len(result) == 1
        assert "good-id" in result[0]["url"]

    def test_multiple_attachments_all_returned(self) -> None:
        """All attachments in the response are returned."""
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "attachments": [
                {"fileInformation": {"fileID": f"id-{i}", "fileName": f"file{i}.pdf"}}
                for i in range(5)
            ]
        }
        with patch("requests.get", return_value=mock_resp):
            result = so.fetch_sam_gov_attachments("NOTICEMULTI")
        assert len(result) == 5
        file_ids = {r["url"].split("/files/")[1].split("/download")[0] for r in result}
        assert file_ids == {f"id-{i}" for i in range(5)}

    def test_attachments_nested_under_data_envelope(self) -> None:
        """Attachments nested under a 'data' key are parsed correctly."""
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "data": {
                "attachments": [
                    {
                        "fileInformation": {
                            "fileID": "nested-id",
                            "fileName": "nested.pdf",
                        }
                    }
                ]
            }
        }
        with patch("requests.get", return_value=mock_resp):
            result = so.fetch_sam_gov_attachments("NOTICE_NESTED")
        assert len(result) == 1
        assert "nested-id" in result[0]["url"]

    def test_noticedesc_endpoint_url_is_queried(self) -> None:
        """The function calls the noticedesc endpoint, not the old search endpoint."""
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {}
        with patch("requests.get", return_value=mock_resp) as mock_get:
            so.fetch_sam_gov_attachments("NOTICE_URL_CHECK", api_key="K2")
        called_url = mock_get.call_args[0][0]
        assert "noticedesc" in called_url
        assert "NOTICE_URL_CHECK" in called_url
        assert "search" not in called_url
